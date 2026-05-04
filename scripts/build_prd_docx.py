"""docs/PRD.md → docs/PRD.docx 的离线转换器。

为什么是这种实现：
  - pandoc 在目标 Windows 环境未安装，需要无外部二进制依赖
  - 直接套 markdown 库再转 docx 会丢失 PRD 里大量"ASCII 流程图代码块"的等宽字符对齐
  - PRD 用的 markdown 子集很小（标题 / 表格 / 代码块 / 列表 / 引用 / 段落 / 分隔线），
    自己写一个行级状态机就够用——比引入额外依赖更可控

支持的 markdown 子集：
  - 标题 # / ## / ### / ####
  - 表格 | header | header |   /   |---|---|
  - 围栏代码块 ```
  - 无序列表 -  /  * （单层）
  - 有序列表 1. （单层）
  - 引用 >
  - 段落
  - 分隔线 --- （Word 里转换成空段落，避免视觉切断）
  - inline: **bold** / *italic* / `inline_code`

故意不支持（PRD 没用到）：
  - 图片 ![]
  - 链接 []()  → 直接输出原文（PRD 里链接量极少）
  - 嵌套列表
  - HTML 块
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Inline 标记的优先级：粗体 → 斜体 → 内联代码
# 用一次性 split 保证不重叠匹配；** 必须先于 * 否则会被吞
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+?`|\*[^*\s][^*]*?\*)")


def _add_runs_with_inline(paragraph, text: str) -> None:
    """把含 inline 标记的字符串拆成多个 run 写到段落里。"""
    if not text:
        return
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            # 浅灰底色提示这是代码（用 shading XML，python-docx 没原生 API）
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F4F4F6")
            r._element.get_or_add_rPr().append(shd)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        else:
            paragraph.add_run(part)


def _set_default_chinese_font(doc: Document, font_name: str = "微软雅黑", size_pt: float = 10.5) -> None:
    """让 Word 在中文环境下默认用雅黑/Yahei；英文区段保持 Calibri 风格。

    python-docx 的 style.font.name 只设西文字体；中文字体必须改 rPr.rFonts 的 eastAsia。
    """
    style = doc.styles["Normal"]
    style.font.size = Pt(size_pt)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def _is_table_separator_row(cells: list[str]) -> bool:
    """| --- | :---: | ---: | 这种第二行视觉分隔。"""
    return all(re.match(r"^:?-+:?$", c.strip()) for c in cells if c.strip())


def _convert(md_path: Path, docx_path: Path) -> None:
    raw = md_path.read_text(encoding="utf-8")
    lines = raw.splitlines()

    doc = Document()
    _set_default_chinese_font(doc)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ---- 空行 ----
        if not stripped:
            i += 1
            continue

        # ---- 分隔线 ----
        if stripped == "---":
            doc.add_paragraph()  # 视觉留白即可，不绘线
            i += 1
            continue

        # ---- 标题 ----
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            h = doc.add_heading(level=level)
            _add_runs_with_inline(h, text)
            i += 1
            continue

        # ---- 围栏代码块 ----
        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # 跳过 closing ```
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            # 紧凑行距，让 ASCII 流程图保持原貌
            pf = p.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)
            pf.line_spacing = 1.05
            continue

        # ---- 表格 ----
        if stripped.startswith("|") and stripped.endswith("|"):
            raw_rows: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                raw_rows.append(lines[i].strip())
                i += 1
            parsed = [
                [c.strip() for c in row.strip("|").split("|")]
                for row in raw_rows
            ]
            # 兼容"有/无"分隔行
            if len(parsed) >= 2 and _is_table_separator_row(parsed[1]):
                header, body = parsed[0], parsed[2:]
            else:
                header, body = parsed[0], parsed[1:]
            ncols = len(header)
            tbl = doc.add_table(rows=1 + len(body), cols=ncols)
            try:
                tbl.style = "Light Grid Accent 1"
            except KeyError:
                tbl.style = "Table Grid"
            # 表头
            for ci, h_txt in enumerate(header):
                cell = tbl.rows[0].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                _add_runs_with_inline(p, h_txt)
                for run in p.runs:
                    run.bold = True
            # 表体
            for ri, row in enumerate(body):
                for ci in range(ncols):
                    txt = row[ci] if ci < len(row) else ""
                    cell = tbl.rows[ri + 1].cells[ci]
                    cell.text = ""
                    _add_runs_with_inline(cell.paragraphs[0], txt)
            doc.add_paragraph()  # 表格后空行
            continue

        # ---- 引用 ----
        if stripped.startswith(">"):
            content = stripped.lstrip(">").strip()
            quote_style = "Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else "Normal"
            p = doc.add_paragraph(style=quote_style)
            _add_runs_with_inline(p, content)
            i += 1
            continue

        # ---- 无序列表 ----
        if re.match(r"^[-*]\s+", stripped):
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                item = re.sub(r"^[-*]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                _add_runs_with_inline(p, item)
                i += 1
            continue

        # ---- 有序列表 ----
        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                _add_runs_with_inline(p, item)
                i += 1
            continue

        # ---- 普通段落 ----
        p = doc.add_paragraph()
        _add_runs_with_inline(p, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main() -> int:
    project_root = Path(__file__).resolve().parent.parent
    md = project_root / "docs" / "PRD.md"
    docx = project_root / "docs" / "PRD.docx"
    if not md.exists():
        print(f"ERROR: PRD.md not found at {md}", file=sys.stderr)
        return 1
    _convert(md, docx)
    size_kb = docx.stat().st_size / 1024
    print(f"OK -> {docx}  ({size_kb:.1f} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
